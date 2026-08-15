# -*- coding: utf-8 -*-
"""将国网新一代协议所有 .doc 转为 .docx，然后全部 .docx 转 .md"""
import os
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import win32com.client
from pathlib import Path

BASE = Path(r"E:\python\南网解析工具\国网新一代协议")

# Step 1: .doc -> .docx via Word COM
print("=" * 60)
print("Step 1: .doc -> .docx 转换")
print("=" * 60)

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = 0
doc_converted = 0
try:
    for f in sorted(BASE.iterdir()):
        if f.suffix == '.doc' and not f.suffixes[-1] == '.docx':
            dst = f.with_suffix('.docx')
            if dst.exists():
                print(f'  已存在，跳过: {f.name}')
                continue
            print(f'  转换: {f.name}')
            doc = word.Documents.Open(str(f), ReadOnly=True)
            doc.SaveAs2(str(dst), FileFormat=16)
            doc.Close(False)
            doc_converted += 1
            print(f'  完成: {dst.name}')
finally:
    word.Quit()

print(f"\n.doc -> .docx 转换完成: {doc_converted} 个文件")

# Step 2: .docx -> .md via python-docx
print("\n" + "=" * 60)
print("Step 2: .docx -> .md 转换")
print("=" * 60)

from docx import Document
from docx.oxml.ns import qn


def build_style_outline_map(doc):
    style_map = {}
    for style in doc.styles:
        if style.type is not None and style.name is not None:
            try:
                rPr = style.element.find(qn('w:rPr'))
                if rPr is not None:
                    outlineLvl = rPr.find(qn('w:outlineLvl'))
                    if outlineLvl is not None:
                        val = outlineLvl.get(qn('w:val'))
                        if val is not None:
                            style_map[style.style_id] = int(val)
            except:
                pass
    return style_map


def para_to_markdown(para, style_outline_map):
    text = para.text.strip()
    if not text:
        return ('', False)
    level = style_outline_map.get(para.style.style_id) if para.style else None
    if level is not None and 0 <= level <= 9:
        return (f'{"#" * (level + 1)} {text}\n', True)
    if para.style and para.style.name and 'List' in para.style.name:
        text = f'- {text}'
    return (text + '\n', False)


def table_to_markdown(table):
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ''
    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append('')
    lines = ['| ' + ' | '.join(rows[0]) + ' |']
    lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')
    for row in rows[1:]:
        lines.append('| ' + ' | '.join(row) + ' |')
    return '\n'.join(lines)


def convert_docx_to_md(docx_path):
    doc = Document(str(docx_path))
    style_map = build_style_outline_map(doc)
    parts = []
    paras = tables = 0
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p':
            for para in doc.paragraphs:
                if para._element is element:
                    text, _ = para_to_markdown(para, style_map)
                    if text:
                        parts.append(text)
                        paras += 1
                    break
        elif tag == 'tbl':
            for table in doc.tables:
                if table._element is element:
                    md = table_to_markdown(table)
                    if md:
                        parts.append('\n' + md + '\n')
                        tables += 1
                    break
    md_path = docx_path.with_suffix('.md')
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(parts))
    return paras, tables


md_converted = 0
total_paras = total_tables = 0
for f in sorted(BASE.glob('*.docx')):
    # Skip temp files
    if f.name.startswith('~$'):
        continue
    md_path = f.with_suffix('.md')
    # Always overwrite for fresh conversion
    print(f'  转换: {f.name}')
    try:
        p, t = convert_docx_to_md(f)
        total_paras += p
        total_tables += t
        md_converted += 1
        md_size = md_path.stat().st_size
        print(f'  完成: {md_path.name} ({p} 段落, {t} 表格, {md_size} bytes)')
    except Exception as e:
        print(f'  错误: {e}')

print(f"\n{'=' * 60}")
print(f"全部完成! {md_converted} 个 .docx -> .md, {total_paras} 段落, {total_tables} 表格")
