"""
Convert 国网新一代协议 docx files to high-precision markdown.
Uses python-docx for structured document conversion.
"""
import os
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


def build_style_outline_map(doc):
    """Build styleId -> outlineLevel mapping from document styles."""
    style_map = {}
    for style in doc.styles:
        if style.type is not None and style.name is not None:
            try:
                rPr = style.element.find(qn('w:rPr'))
                if rPr is not None:
                    outlineLvl = rPr.find(qn('w:outlineLvl'))
                    if outlineLvl is not None:
                        val = outlineLvl.get(qn('w:val'))
                        if val is not None:
                            style_map[style.style_id] = int(val)
            except:
                pass
    return style_map


def get_para_outline_level(para, style_outline_map):
    """Get paragraph outline level from style definition."""
    if para.style and para.style.style_id in style_outline_map:
        return style_outline_map[para.style.style_id]
    return None


def para_to_markdown(para, style_outline_map):
    """Convert paragraph to markdown text. Returns (text, is_heading)."""
    text = para.text.strip()
    if not text:
        return ('', False)
    
    level = get_para_outline_level(para, style_outline_map)
    if level is not None and level >= 0 and level <= 9:
        prefix = '#' * (level + 1)
        return (f'{prefix} {text}\n', True)
    
    # Check if it's a list item
    if para.style and para.style.name and 'List' in para.style.name:
        text = f'- {text}'
    
    return (text + '\n', False)


def table_to_markdown(table):
    """Convert docx table to markdown, handling merged cells."""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace('\n', ' ')
            cells.append(cell_text)
        rows.append(cells)
    
    if not rows:
        return ''
    
    # Build markdown table
    lines = []
    # Header
    lines.append('| ' + ' | '.join(rows[0]) + ' |')
    lines.append('| ' + ' | '.join(['---'] * len(rows[0])) + ' |')
    # Data rows
    for row in rows[1:]:
        # Pad row if needed
        while len(row) < len(rows[0]):
            row.append('')
        lines.append('| ' + ' | '.join(row) + ' |')
    
    return '\n'.join(lines)


def convert_docx_to_md(docx_path, md_path):
    """Convert a single docx file to markdown."""
    doc = Document(docx_path)
    style_outline_map = build_style_outline_map(doc)
    
    md_parts = []
    para_count = 0
    table_count = 0
    
    # Process document body elements in order (paragraphs and tables interleaved)
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        
        if tag == 'p':
            # Paragraph
            for para in doc.paragraphs:
                if para._element is element:
                    text, is_heading = para_to_markdown(para, style_outline_map)
                    if text:
                        md_parts.append(text)
                        para_count += 1
                    break
        elif tag == 'tbl':
            # Table
            for table in doc.tables:
                if table._element is element:
                    md_text = table_to_markdown(table)
                    if md_text:
                        md_parts.append('\n' + md_text + '\n')
                        table_count += 1
                    break
    
    # Write output
    md_content = '\n'.join(md_parts)
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    
    return para_count, table_count


def main():
    base_dir = Path(r"E:\python\南网解析工具\国网新一代协议")
    
    # Find all docx files
    docx_files = sorted(base_dir.glob("*.docx"))
    
    if not docx_files:
        print("No docx files found!")
        return
    
    print(f"Found {len(docx_files)} docx files to convert:")
    
    total_paras = 0
    total_tables = 0
    
    for docx_path in docx_files:
        md_path = docx_path.with_suffix('.md')
        print(f"\nConverting: {docx_path.name}")
        print(f"  -> {md_path.name}")
        
        try:
            paras, tables = convert_docx_to_md(docx_path, md_path)
            total_paras += paras
            total_tables += tables
            
            # Check file size
            md_size = md_path.stat().st_size
            print(f"  OK: {paras} paragraphs, {tables} tables, {md_size} bytes")
        except Exception as e:
            print(f"  ERROR: {e}")
    
    print(f"\n{'='*50}")
    print(f"Conversion complete! {len(docx_files)} files, {total_paras} paragraphs, {total_tables} tables")


if __name__ == '__main__':
    main()
